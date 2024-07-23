from flask import Blueprint, Flask, request, render_template, send_from_directory, redirect, url_for, jsonify, flash
import os
from werkzeug.utils import secure_filename
from datetime import datetime, timedelta
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from flask_mail import Mail, Message
from werkzeug.security import generate_password_hash, check_password_hash
import json
import requests
import msal
import openpyxl
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn as xml_qn
from flask_migrate import Migrate
from flask_bcrypt import Bcrypt

app = Flask(__name__)
bcrypt = Bcrypt(app)

app.config['STATIC_FOLDER'] = 'static'
app.config['UPLOAD_FOLDER'] = 'timesheets'
app.config['PROCESSED_FOLDER'] = 'json/daily'
app.config['ALLOWED_EXTENSIONS'] = {'docx', 'pdf'}

app.config['SECRET_KEY'] = 'your_secret_key'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///site.db'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'docx', 'pdf'}

db = SQLAlchemy(app)
migrate = Migrate(app, db)
login_manager = LoginManager(app)
login_manager.login_view = 'login'
mail = Mail(app)

# Ensure upload and processed directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['PROCESSED_FOLDER'], exist_ok=True)

# Global variable to store config
config = {}

class Organisation(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(150), unique=True, nullable=False)
    logo = db.Column(db.String(150))
    
class Config(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(64), unique=True, nullable=False)
    value = db.Column(db.String(128), nullable=False)

class Role(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(64), unique=True)

class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(64), unique=True, nullable=False)
    password_hash = db.Column(db.String(128), nullable=False)
    role_id = db.Column(db.Integer, db.ForeignKey('role.id'))
    organisation_id = db.Column(db.Integer, db.ForeignKey('organisation.id'))  # Ensure this is defined
    role = db.relationship('Role', backref=db.backref('users', lazy=True))
    organisation = db.relationship('Organisation', backref=db.backref('employees', lazy=True))

    def set_password(self, password):
        self.password_hash = generate_password_hash(password, method='pbkdf2:sha256')

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

    @property
    def is_admin(self):
        return self.role.name == 'admin'

    def __repr__(self):
        return f'<User {self.username}>'

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

class JobCard(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    job_id = db.Column(db.String(50), nullable=False)
    description = db.Column(db.String(200), nullable=False)
    location = db.Column(db.String(100), nullable=False)
    company = db.Column(db.String(100), nullable=False)
    date = db.Column(db.Date, nullable=False)
    price = db.Column(db.Float, nullable=False)
    completed_by = db.Column(db.String(100), nullable=False)
    duration = db.Column(db.Float, nullable=False)
    picture = db.Column(db.String(100))
    video = db.Column(db.String(100))

class Timesheet(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    week_start = db.Column(db.Date, nullable=False)
    date_commencing = db.Column(db.Date, nullable=False)
    hours_worked = db.Column(db.Float, nullable=False)

    user = db.relationship('User', backref=db.backref('timesheets', lazy=True))

# Define the Blueprint for admin routes
admin_bp = Blueprint('admin', __name__)

@admin_bp.route('/admin', methods=['GET', 'POST'])
@login_required
def admin_dashboard():
    if not current_user.is_admin:
        flash('Access denied.', 'danger')
        return redirect(url_for('index'))

    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        role_name = request.form.get('role')
        role = Role.query.filter_by(name=role_name).first()
        if role is None:
            role = Role(name=role_name)
            db.session.add(role)
            db.session.commit()
        user = User(username=username, role=role)
        user.set_password(password)
        db.session.add(user)
        db.session.commit()
        flash('User created successfully.', 'success')

    users = User.query.all()
    return render_template('admin.html', users=users)

@app.route('/account_settings')
@login_required
def account_settings():
    return render_template('account_settings.html')

@app.route('/admin_dashboard')
@login_required
def admin_dashboard():
    if not current_user.is_admin:
        flash('Access denied.', 'danger')
        return redirect(url_for('index'))
    return render_template('admin_dashboard.html')

@admin_bp.route('/create_job_card', methods=['GET', 'POST'])
@login_required
def create_job_card():
    if request.method == 'POST':
        job_id = request.form.get('job_id')
        description = request.form.get('description')
        location = request.form.get('location')
        company = request.form.get('company')
        date = request.form.get('date')
        price = request.form.get('price')
        completed_by = request.form.get('completed_by')
        duration = request.form.get('duration')
        picture = request.files.get('picture')
        video = request.files.get('video')

        picture_filename = secure_filename(picture.filename)
        video_filename = secure_filename(video.filename)
        picture.save(os.path.join('static/uploads', picture_filename))
        video.save(os.path.join('static/uploads', video_filename))

        job_card = JobCard(
            job_id=job_id,
            description=description,
            location=location,
            company=company,
            date=date,
            price=price,
            completed_by=completed_by,
            duration=duration,
            picture=picture_filename,
            video=video_filename
        )
        db.session.add(job_card)
        db.session.commit()
        flash('Job card created successfully.', 'success')
        return redirect(url_for('admin.admin_dashboard'))

    return render_template('create_job_card.html')

@admin_bp.route('/export_job_cards', methods=['GET'])
@login_required
def export_job_cards():
    job_cards = JobCard.query.all()
    data = []
    for job in job_cards:
        data.append({
            'job_id': job.job_id,
            'description': job.description,
            'location': job.location,
            'company': job.company,
            'date': job.date.strftime('%Y-%m-%d'),
            'price': job.price,
            'completed_by': job.completed_by,
            'duration': job.duration
        })
    return jsonify(data)

@admin_bp.route('/create_timesheet', methods=['GET', 'POST'])
@login_required
def create_timesheet():
    if request.method == 'POST':
        week_start = request.form.get('week_start')
        date_commencing = request.form.get('date_commencing')
        hours_worked = request.form.get('hours_worked')
        timesheet = Timesheet(
            user_id=current_user.id,
            week_start=week_start,
            date_commencing=date_commencing,
            hours_worked=hours_worked
        )
        db.session.add(timesheet)
        db.session.commit()
        flash('Timesheet created successfully.', 'success')
        return redirect(url_for('dashboard'))

    return render_template('create_timesheet.html')

@admin_bp.route('/export_timesheets', methods=['GET'])
@login_required
def export_timesheets():
    timesheets = Timesheet.query.filter_by(user_id=current_user.id).all()
    data = []
    for sheet in timesheets:
        data.append({
            'week_start': sheet.week_start.strftime('%Y-%m-%d'),
            'hours_worked': sheet.hours_worked,
            'date_commencing': sheet.date_commencing.strftime('%Y-%m-%d'),
            'user_id': sheet.user_id
        })
    return jsonify(data)

def load_email_config():
    smtp_config = Config.query.filter_by(name='smtp').first()
    if smtp_config:
        app.config['MAIL_SERVER'] = smtp_config.value
        app.config['MAIL_PORT'] = smtp_config.value
        app.config['MAIL_USE_TLS'] = smtp_config.value
        app.config['MAIL_USERNAME'] = smtp_config.value
        app.config['MAIL_PASSWORD'] = smtp_config.value
        app.config['MAIL_DEFAULT_SENDER'] = smtp_config.value
    else:
        print("SMTP configuration not found in the database.")

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def save_timesheet_data(timesheet_data, output_path):
    with open(output_path, 'w') as json_file:
        json.dump(timesheet_data, json_file, indent=4)

@app.route('/', methods=['GET', 'POST'])
@login_required
def index():
    return render_template('index.html', title='Timesheet Processor')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        organisation_id = request.form.get('organisation')
        username = request.form.get('username')
        password = request.form.get('password')
        user = User.query.filter_by(username=username, organisation_id=organisation_id).first()
        if user and user.check_password(password):
            login_user(user)
            return redirect(url_for('index'))
        flash('Login Unsuccessful. Please check username and password', 'danger')
    organisations = Organisation.query.all()
    return render_template('login.html', organisations=organisations)

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('index'))

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        organisation_name = request.form.get('organisation_name')
        username = request.form.get('username')
        password = request.form.get('password')
        organisation = Organisation.query.filter_by(name=organisation_name).first()
        if not organisation:
            organisation = Organisation(name=organisation_name)
            db.session.add(organisation)
            db.session.commit()
        hashed_password = generate_password_hash(password, method='pbkdf2:sha256')
        new_user = User(username=username, password_hash=hashed_password, organisation_id=organisation.id)
        db.session.add(new_user)
        db.session.commit()
        flash('Account created successfully!', 'success')
        return redirect(url_for('login'))
    return render_template('register.html')

@app.route('/dashboard')
@login_required
def dashboard():
    if current_user.role == 'admin':
        timesheets = Timesheet.query.all()
    else:
        timesheets = Timesheet.query.filter_by(user_id=current_user.id).all()
    return render_template('dashboard.html', timesheets=timesheets)

@app.route('/info.html')
def info():
    return render_template('info.html', title='Application Information')

@app.route('/upload', methods=['POST'])
@login_required
def upload_file():
    if 'file' not in request.files:
        return redirect(request.url)
    file = request.files['file']
    if file.filename == '':
        return redirect(request.url)
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        new_timesheet = Timesheet(filename=filename, filepath=filepath, user_id=current_user.id)
        db.session.add(new_timesheet)
        db.session.commit()
        return redirect(url_for('dashboard'))
    return redirect(request.url)

@app.route('/approve/<int:timesheet_id>')
@login_required
def approve_timesheet(timesheet_id):
    if current_user.role != 'admin':
        return redirect(url_for('dashboard'))
    timesheet = Timesheet.query.get_or_404(timesheet_id)
    timesheet.approved = True
    db.session.commit()
    update_excel_file(timesheet.filepath)
    send_approval_email(timesheet.user_id)
    return redirect(url_for('dashboard'))

@app.route('/delete/<int:timesheet_id>')
@login_required
def delete_timesheet(timesheet_id):
    if current_user.role != 'admin':
        return redirect(url_for('dashboard'))
    timesheet = Timesheet.query.get_or_404(timesheet_id)
    db.session.delete(timesheet)
    db.session.commit()
    return redirect(url_for('dashboard'))

def get_access_token(client_id, client_secret, tenant_id, scope):
    authority = f"https://login.microsoftonline.com/{tenant_id}"
    app = msal.ConfidentialClientApplication(
        client_id,
        authority=authority,
        client_credential=client_secret,
    )
    result = app.acquire_token_for_client(scopes=scope)
    if "access_token" in result:
        return result["access_token"]
    else:
        print("Failed to obtain access token")
        print(result.get("error"))
        print(result.get("error_description"))
        print(result.get("correlation_id"))
        return None

def update_excel_file(filepath):
    calculated_json_path = '/mnt/data/employee_data.json'
    with open(calculated_json_path, 'r') as json_file:
        employee_data = json.load(json_file)
    access_token = get_access_token(config['client_id'], config['client_secret'], config['tenant_id'], config['scope'])
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json"
    }
    workbook_url = f"{config['graph_api_endpoint']}/me/drive/root:/{config['excel_file_path']}:/workbook/worksheets"
    response = requests.get(workbook_url, headers=headers)
    if response.status_code != 200:
        print(f"Failed to get worksheet. Status code: {response.status_code}")
        return
    worksheets = response.json()["value"]
    sheet_id = worksheets[0]["id"]
    cell_updates = [
        {
            "range": "A1",
            "values": [[
                "Employee Name", "Week", "Basic", "X1.5", "X2.0", "pay rate/hr", "Annual", "per month"
            ]]
        },
        {
            "range": "A2",
            "values": [[
                employee_data["Employee Name"],
                employee_data["Week"],
                employee_data["Basic"],
                employee_data["X1.5"],
                employee_data["X2.0"],
                employee_data["pay rate/hr"],
                employee_data["Annual"],
                employee_data["per month"]
            ]]
        }
    ]
    for update in cell_updates:
        update_url = f"{config['graph_api_endpoint']}/me/drive/root:/{config['excel_file_path']}:/workbook/worksheets/{sheet_id}/range(address='{update['range']}')"
        response = requests.patch(update_url, headers=headers, json={"values": update["values"]})
        if response.status_code != 200:
            print(f"Failed to update cell {update['range']}. Status code: {response.status_code}")
            return
    print("Excel sheet updated successfully.")
    
def send_approval_email(user_id):
    user = User.query.get(user_id)
    msg = Message("Timesheet Approved", recipients=[user.email])
    msg.body = "Your timesheet has been approved and updated."
    mail.send(msg)

@app.route('/download_template')
def download_template():
    doc = Document()
    current_date = datetime.now()
    week_start = current_date - timedelta(days=current_date.weekday())
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    doc.add_heading('GMT ELECTRICAL SERVICES LTD – WEEKLY TIMESHEET', 0)
    doc.add_paragraph(f'Week Beginning: {week_start.strftime("%d %B %Y")}')
    doc.add_paragraph(f'Name:' f'{"_"*20}')
    table = doc.add_table(rows=7, cols=8, style='Table Grid')
    headers = ["DATE", "WORK SITE ADDRESS", "START", "FINISH", "LUNCH", "BASIC HRS", "O/T 1.5", "O/T 2.0"]
    total_width = Inches(11)
    column_width = total_width / len(headers)
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.width = column_width
        cell.text = header
    for i in range(1, 7):
        row = table.rows[i]
        set_row_height(row, Pt(46))
        if i < 6:
            day = week_start + timedelta(days=i-1)
            row.cells[0].text = day.strftime("%a %d")
    for row in table.rows:
        for cell in row.cells:
            cell.width = column_width
    template_path = os.path.join(app.config['UPLOAD_FOLDER'], 'timesheet_template.docx')
    doc.save(template_path)
    return send_from_directory(app.config['UPLOAD_FOLDER'], 'timesheet_template.docx')

def set_row_height(row, height):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(xml_qn('w:val'), str(height))
    trHeight.set(xml_qn('w:hRule'), "atLeast")
    trPr.append(trHeight)

def load_config():
    global config
    config_data = Config.query.all()
    config = {item.name: item.value for item in config_data}

def populate_config():
    config_data = [
        {'name': 'client_id', 'value': 'your_client_id'},
        {'name': 'client_secret', 'value': 'your_client_secret'},
        {'name': 'tenant_id', 'value': 'your_tenant_id'},
        {'name': 'scope', 'value': 'your_scope'},
        {'name': 'graph_api_endpoint', 'value': 'your_graph_api_endpoint'},
        {'name': 'excel_file_path', 'value': 'your_excel_file_path'},
    ]
    for config_item in config_data:
        existing_config = Config.query.filter_by(name=config_item['name']).first()
        if existing_config is None:
            new_config = Config(name=config_item['name'], value=config_item['value'])
            db.session.add(new_config)
    db.session.commit()

# Register Blueprint
app.register_blueprint(admin_bp, url_prefix='/admin')

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        populate_config()
        load_config()
        load_email_config()
    app.run(debug=True, port=5000)
