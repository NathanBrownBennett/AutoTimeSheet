import os
import json
from tkinter import Tk, filedialog
from docx import Document

def list_files(directory):
    try:
        files = os.listdir(directory)
        files = [f for f in files if os.path.isfile(os.path.join(directory, f))]
        return files
    except FileNotFoundError:
        print(f"The directory {directory} does not exist.")
        return []

def print_menu(files):
    if not files:
        print("No files found.")
        return
    print("Select a file by number:")
    for i, file in enumerate(files, start=1):
        print(f"{i}. {file}")

def get_user_selection(files):
    while True:
        try:
            choice = int(input("Enter the number of the file you want to select: "))
            if 1 <= choice <= len(files):
                return files[choice - 1]
            else:
                print("Invalid number. Please try again.")
        except ValueError:
            print("Invalid input. Please enter a number.")

def open_file_dialog():
    root = Tk()
    root.withdraw()  # Hide the root window
    file_path = filedialog.askopenfilename()
    root.destroy()
    return file_path

def extract_timesheet_data(doc_path):
    doc = Document(doc_path)

    # Initialize dictionary to store extracted data
    timesheet_data = {
        "Employee Name": "",
        "Week Beginning": "",
        "Entries": []
    }

    # Extract the week beginning and employee name from the document
    timesheet_data["Week Beginning"] = doc.paragraphs[-2].text.split("…")[1].strip()
    timesheet_data["Employee Name"] = doc.paragraphs[-1].text.split("…")[1].strip()

    # Extract table data
    table = doc.tables[0]
    for row in table.rows[1:6]:  # Assuming 5 days of data (Mon-Fri)
        cells = row.cells
        entry = {
            "DATE": cells[0].text.strip(),
            "WORK SITE ADDRESS": cells[1].text.strip(),
            "START": cells[2].text.strip(),
            "FINISH": cells[3].text.strip(),
            "LUNCH": cells[4].text.strip(),
            "BASIC HRS": cells[5].text.strip(),
            "O/T 1.5": cells[6].text.strip(),
            "O/T 2.0": cells[7].text.strip()
        }
        timesheet_data["Entries"].append(entry)

    return timesheet_data

def save_timesheet_data(timesheet_data, output_path):
    with open(output_path, 'w') as json_file:
        json.dump(timesheet_data, json_file, indent=4)
    print(f"Timesheet data has been saved to {output_path}")

def calculate_hours_and_pay(json_path):
    with open(json_path, 'r') as json_file:
        timesheet_data = json.load(json_file)

    # Initialize dictionary for calculated data
    employee_data = {
        "Employee Name": timesheet_data["Employee Name"],
        "Week": timesheet_data["Week Beginning"],
        "Basic": 40.0,
        "X1.5": 0.0,
        "X2.0": 0.0,
        "pay rate/hr": 0,  
        "Annual": 0.00,  
        "per month": 0  
    }

    # Calculate total hours
    for entry in timesheet_data["Entries"]:
        employee_data["Basic"] += float(entry["BASIC HRS"])
        employee_data["X1.5"] += float(entry["O/T 1.5"])
        employee_data["X2.0"] += float(entry["O/T 2.0"])

    # Save the calculated data as a JSON file
    calculated_json_path = os.path.join('json/weekly', os.path.basename(json_path).replace('.json', '_calculated.json'))
    with open(calculated_json_path, 'w') as json_file:
        json.dump(employee_data, json_file, indent=4)

    print(f"Employee data has been saved to {calculated_json_path}")
    return calculated_json_path

def update_ignore_list(json_path):
    with open('ignoreList.txt', 'a') as ignore_file:
        ignore_file.write(json_path + '\n')

def load_ignore_list():
    try:
        with open('ignoreList.txt', 'r') as ignore_file:
            ignore_list = ignore_file.read().splitlines()
    except FileNotFoundError:
        ignore_list = []
    return ignore_list

def main():
    fixed_directory = 'timesheets'
    output_directory = 'json/daily'
    os.makedirs(output_directory, exist_ok=True)

    print("Choose an option:")
    print("1. Select a Word document from the fixed directory to extract timesheet data")
    print("2. Open file dialog to select a Word document")
    print("3. Calculate total hours and pay for all JSON files in the output directory")
    print("4. Select a JSON file for calculation from the fixed directory")
    print("5. Open file dialog to select a JSON file for calculation")
    choice = input("Enter your choice (1, 2, 3, 4, or 5): ")

    ignore_list = load_ignore_list()

    if choice == '1':
        files = list_files(fixed_directory)
        if files:
            print_menu(files)
            selected_file = get_user_selection(files)
            selected_file_path = os.path.join(fixed_directory, selected_file)
            print(f"You selected: {selected_file_path}")
        else:
            print("No files available in the directory.")
            return
    elif choice == '2':
        selected_file_path = open_file_dialog()
        if not selected_file_path:
            print("No file selected.")
            return
        print(f"You selected: {selected_file_path}")
    elif choice == '3':
        json_files = list_files(output_directory)
        for json_file in json_files:
            json_file_path = os.path.join(output_directory, json_file)
            if json_file_path not in ignore_list:
                calculated_json_path = calculate_hours_and_pay(json_file_path)
                update_ignore_list(json_file_path)
        return
    elif choice == '4':
        json_files = list_files(output_directory)
        if json_files:
            print_menu(json_files)
            selected_file = get_user_selection(json_files)
            selected_file_path = os.path.join(output_directory, selected_file)
            print(f"You selected: {selected_file_path}")
        else:
            print("No files available in the directory.")
            return
    elif choice == '5':
        selected_file_path = open_file_dialog()
        if not selected_file_path:
            print("No file selected.")
            return
        print(f"You selected: {selected_file_path}")
    else:
        print("Invalid choice.")
        return

    # Extract timesheet data from the selected Word document if choice 1 or 2
    if choice in ['1', '2']:
        timesheet_data = extract_timesheet_data(selected_file_path)

        # Define the output path for the JSON file
        json_filename = os.path.splitext(os.path.basename(selected_file_path))[0] + '.json'
        json_path = os.path.join(output_directory, json_filename)

        # Save the timesheet data as a JSON file
        save_timesheet_data(timesheet_data, 'json/weekly/' + json_filename)

    # Calculate hours and pay if choice 4 or 5
    if choice in ['4', '5']:
        if selected_file_path not in ignore_list:
            calculated_json_path = calculate_hours_and_pay(selected_file_path)
            update_ignore_list(selected_file_path)
        else:
            print("This file has already been processed.")

if __name__ == "__main__":
    main()
