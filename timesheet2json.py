import os
import json
from docx import Document
from tkinter import Tk, filedialog

def list_files(directory):
    try:
        files = os.listdir(directory)
        files = [f for f in files if os.path.isfile(os.path.join(directory, f))]
        return files
    except FileNotFoundError:
        print(f"The directory {directory} does not exist.")
        return []

def print_menu(files):
    if not files:
        print("No files found.")
        return
    print("Select a file by number:")
    for i, file in enumerate(files, start=1):
        print(f"{i}. {file}")

def get_user_selection(files):
    while True:
        try:
            choice = int(input("Enter the number of the file you want to select: "))
            if 1 <= choice <= len(files):
                return files[choice - 1]
            else:
                print("Invalid number. Please try again.")
        except ValueError:
            print("Invalid input. Please enter a number.")

def open_file_dialog():
    root = Tk()
    root.withdraw()  # Hide the root window
    file_path = filedialog.askopenfilename()
    root.destroy()
    return file_path

def extract_timesheet_data(doc_path):
    doc = Document(doc_path)

    # Initialize dictionary to store extracted data
    timesheet_data = {
        "Employee Name": "",
        "Week Beginning": "",
        "Entries": []
    }

    # Function to clean and extract date from text
    def extract_date(text):
        return text.split('…')[-1].strip()

    # Function to find a paragraph that contains a specific keyword, ignoring formatting
    def find_paragraph_with_keyword(doc, keywords):
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            for keyword in keywords:
                if keyword.lower() in text.lower():
                    return text
        return None

    # Search for "Week Beginning" text and extract the date
    week_beginning_text = find_paragraph_with_keyword(doc, ["Week Beginning", "Week Beginning ", "week Beginning:", "Week beginning:"])
    if week_beginning_text:
        timesheet_data["Week Beginning"] = extract_date(week_beginning_text)
    else:
        print("Error: 'Week Beginning' information is not found.")
        return None

    # Search for "Employee Name" or "Name" text and extract the name
    employee_name_text = find_paragraph_with_keyword(doc, ["Employee Name", "Name:", "Name", "Name "])
    if employee_name_text:
        timesheet_data["Employee Name"] = extract_date(employee_name_text)
    else:
        print("Error: 'Employee Name' information is not found.")
        return None

    # Extract table data
    table = doc.tables[0]
    for row in table.rows[1:11]:  # Assuming 7 days of data plus the last two columns (Mon-Sun) + Total normal hours + Total overtime hours
        cells = row.cells
        entry = {
            "DATE": cells[0].text.strip(),
            "WORK SITE ADDRESS": cells[1].text.strip(),
            "START": cells[2].text.strip(),
            "FINISH": cells[3].text.strip(),
            "LUNCH": cells[4].text.strip(),
            "BASIC HRS": cells[5].text.strip(),
            "O/T 1.5": cells[6].text.strip(),
            "O/T 2.0": cells[7].text.strip()
        }
        timesheet_data["Entries"].append(entry)

    return timesheet_data

def save_timesheet_data(timesheet_data, output_path):
    with open(output_path, 'w') as json_file:
        json.dump(timesheet_data, json_file, indent=4)
    print(f"Timesheet data has been saved to {output_path}")

def main():
    fixed_directory = 'timesheets'
    output_directory = 'json/daily'
    os.makedirs(output_directory, exist_ok=True)

    print("Choose an option:")
    print("1. Select a file from the fixed directory")
    print("2. Open file dialog to select a file")
    choice = input("Enter your choice (1 or 2): ")

    if choice == '1':
        files = list_files(fixed_directory)
        if files:
            print_menu(files)
            selected_file = get_user_selection(files)
            selected_file_path = os.path.join(fixed_directory, selected_file)
            print(f"You selected: {selected_file_path}")
        else:
            print("No files available in the directory.")
            return
    elif choice == '2':
        selected_file_path = open_file_dialog()
        if not selected_file_path:
            print("No file selected.")
            return
        print(f"You selected: {selected_file_path}")
    else:
        print("Invalid choice.")
        return

    # Extract timesheet data from the selected file
    timesheet_data = extract_timesheet_data(selected_file_path)

    # Define the output path for the JSON file
    json_filename = os.path.splitext(os.path.basename(selected_file_path))[0] + '.json'
    json_path = os.path.join(output_directory, json_filename)

    # Save the timesheet data as a JSON file
    save_timesheet_data(timesheet_data, json_path)

if __name__ == "__main__":
    main()
